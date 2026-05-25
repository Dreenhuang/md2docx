"""
DOCX 生成器模块
将结构化 DocumentNode 树转换为符合 PRD 规范的 Word 文档

格式规范（严格遵守）：
- 纸张 A4 纵向 (210mm x 297mm)
- 页边距：左20mm / 右20mm / 上16mm / 下16mm
- 全局字体：微软雅黑纯黑色
- 行距：固定值 16 磅
- 正文首行缩进：2 字符（约 742080 EMU）
- 正文字号：小四 12pt 常规体
- H1=26pt, H2=22pt, H3=18pt, H4=16pt, H5=14pt, H6=13pt（全部 Bold + 较正文加大一号）
- 加粗重点内容：Bold + 14pt + 纯黑
"""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class DocxGenerator:
    """
    DOCX 文档生成器

    将 MarkdownParser 输出的 DocumentNode 列表转换为 .docx 文件
    所有格式参数严格遵循 PRD 文档规范
    """

    FONT_FAMILY = "微软雅黑"
    FONT_COLOR = RGBColor(0, 0, 0)

    FONT_SIZE_BODY = Pt(12)

    HEADING_SIZES = {
        1: Pt(26),
        2: Pt(22),
        3: Pt(18),
        4: Pt(16),
        5: Pt(14),
        6: Pt(13),
    }

    LINE_SPACING = 16
    FIRST_LINE_INDENT = Emu(742080)

    CODE_FONT = "Consolas"
    CODE_BG_COLOR = "F0F0F0"

    def __init__(self, config: dict = None):
        """
        初始化生成器

        参数：
            config: 配置字典，可覆盖默认值
                - paper_size: 纸张大小（默认 A4）
                - margin_left/right/top/bottom: 页边距 mm
                - font_family: 字体名称
                - line_spacing: 行距磅值
                - first_line_indent: 首行缩进字符数
        """
        self.config = config or {}
        self._font_family = self.config.get('font_family', self.FONT_FAMILY)
        self._line_spacing = float(self.config.get('line_spacing', self.LINE_SPACING))
        indent_chars = int(self.config.get('first_line_indent', 2))
        self._first_line_indent = Emu(742080 * (indent_chars // 2 or 1))

    def generate(self, nodes: list, doc_title: str = "") -> Document:
        """
        生成 DOCX 文档对象

        参数：
            nodes: DocumentNode 节点列表（来自 MarkdownParser）
            doc_title: 文档标题（用于元信息区）

        返回：
            docx.Document 对象
        """
        doc = Document()
        self._setup_page(doc)
        self._render_nodes(doc, nodes)
        return doc

    def save(self, document: Document, output_path: str) -> str:
        """
        保存文档到指定路径

        参数：
            document: DocxDocument 对象
            output_path: 输出文件完整路径

        返回：
            实际保存的文件路径
        """
        dir_path = os.path.dirname(output_path)
        if dir_path and not os.path.exists(dir_path):
            os.makedirs(dir_path, exist_ok=True)
        document.save(output_path)
        return output_path

    def _setup_page(self, doc: Document):
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        margin_left_mm = float(self.config.get('margin_left', 20))
        margin_right_mm = float(self.config.get('margin_right', 20))
        margin_top_mm = float(self.config.get('margin_top', 16))
        margin_bottom_mm = float(self.config.get('margin_bottom', 16))
        section.left_margin = Cm(margin_left_mm / 10.0)
        section.right_margin = Cm(margin_right_mm / 10.0)
        section.top_margin = Cm(margin_top_mm / 10.0)
        section.bottom_margin = Cm(margin_bottom_mm / 10.0)
        section.header_distance = Cm(0.85)
        section.footer_distance = Cm(0.85)

    def _render_nodes(self, doc: Document, nodes: list):
        """遍历节点列表，逐个渲染到文档中"""
        for node in nodes:
            self._render_node(doc, node)

    def _render_node(self, doc: Document, node):
        """根据节点类型分发渲染方法"""
        dispatch = {
            'heading': self._render_heading,
            'paragraph': self._render_paragraph,
            'blockquote': self._render_blockquote,
            'table': self._render_table,
            'unordered_list': self._render_list,
            'ordered_list': self._render_list,
            'code_block': self._render_code_block,
            'image': self._render_image,
            'horizontal_rule': self._render_horizontal_rule,
        }
        handler = dispatch.get(node.node_type, self._render_paragraph)
        handler(doc, node)

    def _render_heading(self, doc: Document, node):
        """
        渲染标题节点

        H1-H6 全部 Bold + 加大字号 + 纯黑
        H1=26pt, H2=22pt, H3=18pt, H4=16pt, H5=14pt, H6=13pt
        """
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._apply_base_format(para, first_indent=False)

        level = max(1, min(6, node.level))
        size = self.HEADING_SIZES.get(level, Pt(14))
        run = para.add_run(node.content)
        run.font.name = self._font_family
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self._font_family)
        run.font.size = size
        run.font.bold = True
        run.font.color.rgb = self.FONT_COLOR

        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)

    def _render_paragraph(self, doc: Document, node):
        """
        渲染普通段落

        正文字号 12pt 常规体微软雅黑
        检测 **加粗** 内容并特殊渲染
        """
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self._apply_base_format(para, first_indent=True)

        bold_parts = node.metadata.get('bold_parts', [])
        if bold_parts:
            self._render_rich_text(para, node.content, bold_parts)
        else:
            run = para.add_run(node.content)
            self._set_body_style(run)

    def _render_rich_text(self, para, text: str, bold_parts: list):
        """
        渲染包含加粗标记的富文本

        将 **text** 部分用 Bold + 14pt 渲染
        其余部分用正文字号 12pt 渲染
        """
        pattern = r'(\*\*.*?\*\*)'
        parts = re.split(pattern, text)

        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                clean_text = part[2:-2]
                run = para.add_run(clean_text)
                run.font.name = self._font_family
                run._element.rPr.rFonts.set(qn('w:eastAsia'), self._font_family)
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = self.FONT_COLOR
            else:
                run = para.add_run(part)
                self._set_body_style(run)

    def _render_blockquote(self, doc: Document, node):
        """
        渲染引用块节点

        左缩进 + 左边框线 + 斜体文字
        """
        para = doc.add_paragraph()
        self._apply_base_format(para, first_indent=False)
        para.paragraph_format.left_indent = Cm(1.0)
        lines = node.content.split('\n')
        for i, line in enumerate(lines):
            if i > 0:
                run = para.add_run('\n')
            run = para.add_run(line)
            run.font.name = self._font_family
            run.font.size = self.FONT_SIZE_BODY
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            pPr = para._p.get_or_add_pPr()
            border = OxmlElement('w:pBdr')
            left_border = OxmlElement('w:left')
            left_border.set(qn('w:val'), 'single')
            left_border.set(qn('w:sz'), '12')
            left_border.set(qn('w:color'), '999999')
            border.append(left_border)
            pPr.append(border)

    def _render_table(self, doc: Document, node):
        """
        渲染表格节点

        单线边框、文字纯黑、表头加粗
        """
        headers = node.metadata.get('headers', [])
        rows = [child.metadata.get('cells', []) for child in node.children
                if child.node_type == 'table_row']

        total_cols = len(headers) if headers else (len(rows[0]) if rows else 1)
        table = doc.add_table(rows=1 + len(rows), cols=total_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, header in enumerate(headers):
            if i < total_cols:
                cell = table.rows[0].cells[i]
                cell.text = ''
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(str(header))
                run.font.name = self._font_family
                run._element.rPr.rFonts.set(qn('w:eastAsia'), self._font_family)
                run.font.size = Pt(11)
                run.font.bold = True
                run.font.color.rgb = self.FONT_COLOR
                self._set_cell_border(cell, single_color="000000")

        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < total_cols:
                    cell = table.rows[row_idx + 1].cells[col_idx]
                    cell.text = ''
                    p = cell.paragraphs[0]
                    run = p.add_run(str(cell_text))
                    run.font.name = self._font_family
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), self._font_family)
                    run.font.size = Pt(10.5)
                    run.font.color.rgb = self.FONT_COLOR
                    self._set_cell_border(cell, single_color="000000")

    def _set_cell_border(self, cell, single_color="000000"):
        """设置单元格单线边框"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:color'), single_color)
            borders.append(border)
        tcPr.append(borders)

    def _render_list(self, doc: Document, node):
        """
        渲染有序/无序列表

        使用段落模拟列表效果，保持格式一致性
        """
        is_ordered = node.node_type == 'ordered_list'
        for idx, child in enumerate(node.children):
            para = doc.add_paragraph()
            self._apply_base_format(para, first_indent=True)

            if is_ordered:
                prefix = f"{idx + 1}. "
            else:
                prefix = "• "

            run_prefix = para.add_run(prefix)
            self._set_body_style(run_prefix)

            run_content = para.add_run(child.content)
            self._set_body_style(run_content)

    def _render_code_block(self, doc: Document, node):
        """
        渲染代码块

        Consolas 等宽字体 + 浅灰背景
        保持原始缩进和换行
        """
        para = doc.add_paragraph()
        self._apply_base_format(para, first_indent=False)

        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), self.CODE_BG_COLOR)
        para._p.get_or_add_pPr().append(shading_elm)

        lines = node.content.split('\n')
        full_text = '\n'.join(lines)
        run = para.add_run(full_text)
        run.font.name = self.CODE_FONT
        run.font.size = Pt(9)
        run.font.color.rgb = self.FONT_COLOR

    def _render_image(self, doc: Document, node):
        """
        渲染图片节点

        尝试嵌入本地图片文件
        如果图片不存在则显示占位文字
        """
        src = node.metadata.get('src', '')
        alt_text = node.content or '图片'

        if src and os.path.exists(src):
            try:
                width = Cm(14)
                last_para = doc.paragraphs[-1] if doc.paragraphs else None
                run = (last_para.add_run() if last_para else doc.add_paragraph().add_run())
                run.add_picture(src, width=width)
                return
            except Exception:
                pass

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._apply_base_format(para, first_indent=False)
        run = para.add_run(f"[图片: {alt_text}]")
        self._set_body_style(run)

    def _render_horizontal_rule(self, doc: Document, node):
        """渲染分隔线（空行代替）"""
        doc.add_paragraph()

    def _apply_base_format(self, para, first_indent: bool = True):
        """
        应用基础段落格式

        - 行距最小值 16 磅（AT_LEAST 模式，大字号自动扩展）
        - 首行缩进 2 字符（可选）
        - 段前段后间距 0
        """
        para.paragraph_format.line_spacing = Pt(self._line_spacing)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)

        if first_indent:
            para.paragraph_format.first_line_indent = self._first_line_indent
        else:
            para.paragraph_format.first_line_indent = Emu(0)

    def _set_body_style(self, run):
        """
        设置正文 run 样式

        微软雅黑 / 小四 12pt / 常规体 / 纯黑色
        """
        run.font.name = self._font_family
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self._font_family)
        run.font.size = self.FONT_SIZE_BODY
        run.font.color.rgb = self.FONT_COLOR
