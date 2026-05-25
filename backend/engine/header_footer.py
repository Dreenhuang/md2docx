"""
页眉页脚生成器模块
为 DOCX 文档添加符合 PRD 规范的页眉和页脚

规范：
- 页眉：文档标题居中 / 微软雅黑 9pt / 纯黑
- 页脚："第 X 页 / 共 Y 页 ｜ 编写日期：YYYY年MM月DD日" / 居中 / 9pt 微软雅黑 / 纯黑
- 标题提取优先级：YAML Front Matter title > 第一个 H1 > 文件名(去扩展名)
- 页码从正文第一页开始编号
"""

from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class HeaderFooterGenerator:
    """
    页眉页脚生成器

    为 Word 文档添加标准化的页眉（标题）和页脚（页码+日期）
    """

    FONT_FAMILY = "微软雅黑"
    FONT_SIZE = Pt(9)
    FONT_COLOR = RGBColor(0, 0, 0)

    @staticmethod
    def apply(document: Document, title_str: str, date_str: str = None) -> Document:
        """
        为文档应用页眉和页脚

        参数：
            document: docx.Document 对象
            title_str: 页眉显示的文档标题
            date_str: 编写日期字符串，格式 YYYY年MM月DD日。默认使用当前日期

        返回：
            同一 document 对象（支持链式调用）
        """
        if not date_str:
            date_str = datetime.now().strftime("%Y年%m月%d日")

        generator = HeaderFooterGenerator()
        generator._add_header(document, title_str)
        generator._add_footer(document, date_str)

        return document

    @staticmethod
    def extract_title(front_matter: dict, nodes: list, file_name: str) -> str:
        """
        按优先级提取文档标题

        优先级：
        1. YAML Front Matter 中的 title 字段
        2. 第一个 H1 标题的内容
        3. 文件名（去掉扩展名）

        参数：
            front_matter: YAML 前置元数据字典
            nodes: 解析后的节点列表
            file_name: 原始文件名

        返回：
            提取到的标题字符串
        """
        if front_matter and front_matter.get('title'):
            return str(front_matter['title'])

        for node in nodes:
            if node.node_type == 'heading' and node.level == 1:
                return node.content

        if file_name:
            base_name = file_name.rsplit('.', 1)[0] if '.' in file_name else file_name
            return base_name

        return "未命名文档"

    def _add_header(self, document: Document, title_text: str):
        """
        添加页眉：文档标题居中 / 微软雅黑 9pt / 纯黑
        """
        section = document.sections[0]
        header = section.header
        header.is_linked_to_previous = False

        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.text = ''

        run = para.add_run(title_text)
        run.font.name = self.FONT_FAMILY
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.FONT_FAMILY)
        run.font.size = self.FONT_SIZE
        run.font.color.rgb = self.FONT_COLOR

    def _add_footer(self, document: Document, date_str: str):
        """
        添加页脚："第 X 页 / 共 Y 页 ｜ 编写日期：YYYY年MM月DD日"
        居中 / 小五号 9pt 微软雅黑 / 纯黑
        使用 PAGE 和 NUMPAGES 域实现动态页码
        """
        section = document.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False

        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.text = ''

        run1 = para.add_run("第 ")
        self._set_footer_style(run1)

        page_field = self._create_page_number_field(para)
        para._p.append(page_field)

        run2 = para.add_run(" 页 / 共 ")
        self._set_footer_style(run2)

        total_field = self._create_total_pages_field(para)
        para._p.append(total_field)

        run3 = para.add_run(f" 页 ｜ 编写日期：{date_str}")
        self._set_footer_style(run3)

    def _create_page_number_field(self, para):
        """创建当前页码域 (PAGE)"""
        run = OxmlElement('w:r')
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run.append(fldChar1)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = " PAGE "
        run.append(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run.append(fldChar2)

        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run.append(fldChar3)

        return run

    def _create_total_pages_field(self, para):
        """创建总页码域 (NUMPAGES)"""
        run = OxmlElement('w:r')
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run.append(fldChar1)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = " NUMPAGES "
        run.append(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run.append(fldChar2)

        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run.append(fldChar3)

        return run

    def _set_footer_style(self, run):
        """设置页脚文字样式"""
        run.font.name = self.FONT_FAMILY
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.FONT_FAMILY)
        run.font.size = self.FONT_SIZE
        run.font.color.rgb = self.FONT_COLOR
